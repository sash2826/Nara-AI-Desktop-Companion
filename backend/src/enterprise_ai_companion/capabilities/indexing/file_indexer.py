"""Indexes local workspace files into SQLite, Qdrant, and the knowledge graph."""

from __future__ import annotations

import asyncio
import hashlib
import logging
import platform
import uuid
from collections.abc import Callable
from contextlib import asynccontextmanager
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from weakref import WeakValueDictionary

from enterprise_ai_companion.capabilities.graph.graph_provider import GraphProvider
from enterprise_ai_companion.capabilities.graph.graph_state_repository import GraphStateRepository
from enterprise_ai_companion.capabilities.graph.knowledge_graph_service import KnowledgeGraphService
from enterprise_ai_companion.capabilities.graph.null_graph_provider import NullGraphProvider
from enterprise_ai_companion.capabilities.indexing.chunk_repository import (
    Chunk,
    ChunkRepository,
)
from enterprise_ai_companion.capabilities.indexing.document_repository import (
    DocumentRepository,
    IndexedDocument,
)
from enterprise_ai_companion.capabilities.indexing.embedding_service import EmbeddingService
from enterprise_ai_companion.capabilities.indexing.abbreviation_repository import (
    AbbreviationRepository,
)
from enterprise_ai_companion.capabilities.indexing.indexing_error_repository import (
    IndexingErrorRepository,
)
from enterprise_ai_companion.capabilities.indexing.text_chunker import TextChunker
from enterprise_ai_companion.capabilities.retrieval.abbreviation_extractor import (
    AbbreviationExtractor,
)

# TYPE_CHECKING guard avoids a circular import at runtime; PluginManager
# is only needed as a type annotation here.
from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from enterprise_ai_companion.capabilities.plugins.plugin_manager import PluginManager

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".xlsx"}

# Maximum number of files processed concurrently during workspace indexing.
# Kept at 1 because all repositories share a single aiosqlite connection; with
# concurrent coroutines the transaction-state tracking (in_transaction flag)
# desynchronises — one coroutine's COMMIT unexpectedly finalises another's
# in-flight writes, leading to "cannot start a transaction within a transaction"
# and SQLITE_MISUSE errors under watchdog-triggered multi-file drops.
_INDEXING_CONCURRENCY = 1

# Directory names that are never traversed during indexing, regardless of depth.
# Shared with file_watcher.py which imports this constant.
EXCLUDED_DIRS: frozenset[str] = frozenset({
    "node_modules",
    ".git",
    ".venv",
    ".claude",
    "__pycache__",
    "target",          # Rust build output
    "dist",
    "build",
    ".next",
    ".nuxt",
    "$Recycle.Bin",
    "Windows",
    "Program Files",
    "Program Files (x86)",
    "AppData",
    "System Volume Information",
})

# FILE_ATTRIBUTE_RECALL_ON_DATA_ACCESS (0x400000) — Windows sets this on
# OneDrive Files On-Demand stubs that are not yet downloaded locally.
_ONEDRIVE_STUB_ATTR = 0x400000

# OS-critical roots that must never be indexed, regardless of what path the
# caller supplies. Prevents a confused-deputy attack where a malicious document
# or UI action triggers indexing of system directories.
_BLOCKED_ROOTS: frozenset[Path] = frozenset({
    # Windows system directories
    Path("C:/Windows"),
    Path("C:/Program Files"),
    Path("C:/Program Files (x86)"),
    Path("C:/ProgramData"),
    Path("C:/System Volume Information"),
    # POSIX system directories
    Path("/etc"),
    Path("/sys"),
    Path("/proc"),
    Path("/dev"),
    Path("/boot"),
})


def _collect_files(root: Path) -> list[Path]:
    """Walk root recursively, skipping any directory in EXCLUDED_DIRS.

    Using os.walk instead of Path.rglob so we can prune entire subtrees
    (e.g. node_modules) without descending into them first.
    """
    import os

    collected: list[Path] = []
    for dirpath, dirnames, filenames in os.walk(root):
        # Prune excluded dirs in-place so os.walk does not recurse into them.
        dirnames[:] = [d for d in dirnames if d not in EXCLUDED_DIRS]
        for name in filenames:
            p = Path(dirpath) / name
            if p.suffix.lower() in SUPPORTED_EXTENSIONS:
                collected.append(p)
    return collected


def _is_cloud_stub(path: Path) -> bool:
    """Return True if the file is a OneDrive cloud-only placeholder.

    On non-Windows platforms this always returns False.
    """
    if platform.system() != "Windows":
        return False
    try:
        import ctypes
        attrs = ctypes.windll.kernel32.GetFileAttributesW(str(path))
        return attrs != -1 and bool(attrs & _ONEDRIVE_STUB_ATTR)
    except Exception:
        return False


@dataclass
class IndexingResult:
    files_found: int = 0
    files_indexed: int = 0
    files_skipped: int = 0
    errors: list[str] = field(default_factory=list)

    @property
    def status(self) -> str:
        return "completed" if not self.errors else "completed_with_errors"


class FileIndexer:
    """Recursively indexes text files in a workspace directory.

    For each file that is new or has changed (detected by SHA-256 hash), the
    indexer chunks the content, generates embeddings, persists both to SQLite
    (via DocumentRepository / ChunkRepository) and Qdrant, then builds graph
    entities via KnowledgeGraphService (best-effort — graph failures never abort
    indexing).
    """

    def __init__(
        self,
        doc_repo: DocumentRepository,
        chunk_repo: ChunkRepository,
        embedding_service: EmbeddingService,
        chunker: TextChunker | None = None,
        graph_provider: GraphProvider | None = None,
        error_repo: IndexingErrorRepository | None = None,
        abbreviation_extractor: AbbreviationExtractor | None = None,
        abbreviation_repo: AbbreviationRepository | None = None,
        graph_state_repo: GraphStateRepository | None = None,
        plugin_manager: "PluginManager | None" = None,
    ) -> None:
        self._doc_repo = doc_repo
        self._chunk_repo = chunk_repo
        self._embedding_service = embedding_service
        self._chunker = chunker or TextChunker()
        self._graph_service = KnowledgeGraphService(graph_provider or NullGraphProvider())
        self._error_repo = error_repo
        self._abbreviation_extractor = abbreviation_extractor
        self._abbreviation_repo = abbreviation_repo
        self._graph_state_repo = graph_state_repo
        self._plugin_manager = plugin_manager
        # Per-file locks prevent concurrent indexing of the same path, which
        # causes FOREIGN KEY failures when one coroutine replaces the document
        # row while another is mid-way through inserting its graph entities.
        # WeakValueDictionary releases locks automatically when no coroutine
        # holds a reference, so the dict never grows unbounded.
        self._file_locks: WeakValueDictionary[str, asyncio.Lock] = WeakValueDictionary()
        self._locks_meta = asyncio.Lock()
        # Per-workspace locks prevent two watcher-triggered index_workspace calls
        # for the same folder from running concurrently. Even with
        # _INDEXING_CONCURRENCY=1, two concurrent calls would still interleave
        # their single-file sequences and race on the shared DB connection.
        self._workspace_locks: dict[str, asyncio.Lock] = {}
        self._workspace_locks_meta = asyncio.Lock()
        # Strong references to fire-and-forget background tasks (Pass 2 graph
        # extraction). asyncio only holds weak refs; without this the GC can
        # collect a running task before it finishes.
        self._background_tasks: set[asyncio.Task] = set()

    def _is_safe_path(self, resolved: Path) -> bool:
        """Return False if resolved is inside a blocked OS-critical directory."""
        for root in _BLOCKED_ROOTS:
            try:
                resolved.relative_to(root.resolve())
                return False
            except ValueError:
                pass
        return True

    @asynccontextmanager
    async def _file_lock(self, path: str):
        """Yield a per-file asyncio.Lock, creating it if it does not yet exist.

        Serialises concurrent _index_file calls for the same path so that one
        coroutine cannot replace the document row (invalidating its UUID) while
        another coroutine is mid-way through inserting graph entities for the
        old UUID, which would produce a FOREIGN KEY constraint failure.

        WeakValueDictionary means the Lock is released from memory as soon as
        no coroutine holds a live reference, so the dict never grows unbounded.
        """
        async with self._locks_meta:
            lock = self._file_locks.get(path)
            if lock is None:
                lock = asyncio.Lock()
                self._file_locks[path] = lock
        async with lock:
            yield

    @asynccontextmanager
    async def _workspace_lock(self, path: str):
        """Yield a per-workspace asyncio.Lock, serialising concurrent index_workspace
        calls for the same folder (e.g. two rapid watcher events)."""
        async with self._workspace_locks_meta:
            lock = self._workspace_locks.get(path)
            if lock is None:
                lock = asyncio.Lock()
                self._workspace_locks[path] = lock
        async with lock:
            yield

    async def index_workspace(
        self,
        workspace_path: str,
        progress_cb: Callable[[IndexingResult], None] | None = None,
    ) -> IndexingResult:
        """Index all supported files under workspace_path. Returns a summary.

        progress_cb is called after every file is processed so callers can track
        live progress rather than waiting for the full run to finish.
        """
        async with self._workspace_lock(workspace_path):
            return await self._index_workspace_locked(workspace_path, progress_cb)

    async def _index_workspace_locked(
        self,
        workspace_path: str,
        progress_cb: Callable[[IndexingResult], None] | None = None,
    ) -> IndexingResult:
        root = Path(workspace_path).resolve()

        if not self._is_safe_path(root):
            raise ValueError(f"Unsafe workspace path rejected: {root}")

        result = IndexingResult()

        if not root.exists() or not root.is_dir():
            result.errors.append(f"Workspace path does not exist or is not a directory: {root}")
            return result

        files = _collect_files(root)
        result.files_found = len(files)

        # Pass 1 — text extraction + chunking + embeddings + SQLite + Qdrant.
        # Up to _INDEXING_CONCURRENCY files run concurrently; the expensive parts
        # (OCR, ONNX inference, Qdrant writes) all run in thread-pool executors
        # so the event loop stays free throughout.
        deferred_graph: list[tuple[str, str, list[str], str]] = []
        sem = asyncio.Semaphore(_INDEXING_CONCURRENCY)

        async def _process(fp: Path) -> None:
            # Yield first so CancelledError is delivered promptly between files.
            await asyncio.sleep(0)
            if _is_cloud_stub(fp):
                logger.debug("Skipping cloud-only stub: %s", fp.name)
                result.files_skipped += 1
                if progress_cb:
                    progress_cb(result)
                return
            try:
                indexed = await self._index_file(fp, workspace_path, deferred_graph)
                if indexed:
                    result.files_indexed += 1
                else:
                    result.files_skipped += 1
            except Exception as exc:
                logger.warning("Failed to index %s: %s", fp, exc)
                result.errors.append(f"{fp}: {exc}")
                if self._error_repo is not None:
                    try:
                        await self._error_repo.save(workspace_path, str(fp), str(exc))
                    except Exception as repo_exc:
                        logger.warning("Failed to persist indexing error: %s", repo_exc)
            if progress_cb:
                progress_cb(result)

        async def _bounded(fp: Path) -> None:
            async with sem:
                await _process(fp)

        await asyncio.gather(*(_bounded(fp) for fp in files))

        # Pass 2 — graph entity extraction via LLM runs as a background task so
        # search is available immediately after Pass 1 completes.
        if deferred_graph:
            logger.info("[PASS 2] Scheduling graph extraction for %d document(s)", len(deferred_graph))
            task = asyncio.create_task(self._run_graph_pass(deferred_graph))
            self._background_tasks.add(task)
            task.add_done_callback(self._background_tasks.discard)

        logger.info(
            "Indexing complete: %d found, %d indexed, %d skipped, %d errors",
            result.files_found,
            result.files_indexed,
            result.files_skipped,
            len(result.errors),
        )
        return result

    async def index_file(self, file_path: str, workspace_path: str) -> str | None:
        """Index a single file and return its document ID on success.

        Returns the document ID whether the file was freshly indexed or was
        already up-to-date (unchanged hash). Returns None only when the file
        is unsupported, unsafe, or caused an unrecoverable error.

        Returning the ID for unchanged files ensures callers (e.g. the Downloads
        watcher hook) can still act on already-indexed files dropped into a new
        location — the file is present in the index, just with no content change.
        """
        resolved = Path(file_path).resolve()
        if not self._is_safe_path(resolved):
            logger.warning("index_file: unsafe path rejected: %s", resolved)
            return None

        try:
            await self._index_file(resolved, workspace_path)
        except Exception as exc:
            logger.warning("index_file: error indexing %s: %s", resolved.name, exc)
            return None

        doc = await self._doc_repo.get_by_path(str(resolved))
        return doc.id if doc else None

    async def delete_file(self, file_path: str) -> bool:
        """Remove all index records for a file. Returns True if the file was indexed.

        Cleans up chunks (SQLite + Qdrant), graph entities, graph state, and the
        document row. Safe to call when the file no longer exists on disk.
        """
        resolved = str(Path(file_path).resolve())
        doc = await self._doc_repo.get_by_path(resolved)
        if doc is None:
            return False

        await self._chunk_repo.delete_by_document(doc.id)
        await self._graph_service.delete_document(doc.id)
        if self._graph_state_repo is not None:
            await self._graph_state_repo.delete_by_document(doc.id)
        await self._doc_repo.delete_by_path(resolved)

        logger.debug("Deleted index for %s", Path(resolved).name)
        return True

    async def rename_file(self, src_path: str, dest_path: str) -> bool:
        """Update the stored file_path when a file is renamed or moved within a workspace.

        Only updates the SQLite path — no re-embedding or re-chunking. Returns
        True when the source document was found and updated.
        """
        src = str(Path(src_path).resolve())
        dest = str(Path(dest_path).resolve())
        doc = await self._doc_repo.get_by_path(src)
        if doc is None:
            return False

        await self._doc_repo.update_path(src, dest)
        logger.debug("Renamed index: %s → %s", Path(src).name, Path(dest).name)
        return True

    def _extract_text(self, file_path: Path) -> str:
        """Extract plain text from a file based on its extension.

        Enabled FileProcessorPlugins are consulted first so they can handle
        custom extensions or override built-in extraction for known types.
        """
        ext = file_path.suffix.lower()

        if self._plugin_manager:
            for fp_plugin in self._plugin_manager.get_file_processors():
                if ext in fp_plugin.supported_extensions:
                    return fp_plugin.extract_text(file_path)

        if ext in {".txt", ".md"}:
            return file_path.read_text(encoding="utf-8", errors="replace")
        if ext == ".pdf":
            import pypdf  # noqa: PLC0415
            reader = pypdf.PdfReader(str(file_path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        if ext == ".docx":
            import docx  # noqa: PLC0415
            doc = docx.Document(str(file_path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if ext == ".xlsx":
            import openpyxl  # noqa: PLC0415
            wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
            parts: list[str] = []
            for sheet in wb.worksheets:
                if sheet.title:
                    parts.append(sheet.title)
                for row in sheet.iter_rows(values_only=True):
                    row_text = "  ".join(str(v) for v in row if v is not None)
                    if row_text.strip():
                        parts.append(row_text)
            wb.close()
            return "\n".join(parts)
        return file_path.read_text(encoding="utf-8", errors="replace")

    async def _index_file(
        self,
        file_path: Path,
        workspace_path: str,
        deferred: list | None = None,
    ) -> bool:
        """Index a single file. Returns True if indexed, False if unchanged.

        Acquires a per-file lock before doing any database work so that two
        concurrent calls for the same path are serialised.  Without this, one
        coroutine can delete and replace the document row while another is
        inserting graph entities for the old UUID, causing FOREIGN KEY failures.

        deferred — when provided, graph extraction is appended to this list
        instead of being run inline. Used by index_workspace (Pass 2 pattern).
        """
        async with self._file_lock(str(file_path)):
            return await self._index_file_locked(file_path, workspace_path, deferred)

    async def _index_file_locked(
        self,
        file_path: Path,
        workspace_path: str,
        deferred: list | None = None,
    ) -> bool:
        """Inner implementation — called only while the per-file lock is held."""
        loop = asyncio.get_running_loop()
        text = await loop.run_in_executor(None, self._extract_text, file_path)

        # Run text through enabled TextProcessorPlugins before hashing/chunking.
        if self._plugin_manager:
            for tp_plugin in self._plugin_manager.get_text_processors():
                try:
                    text = tp_plugin.process_text(text, str(file_path))
                except Exception as exc:  # noqa: BLE001
                    logger.warning(
                        "TextProcessorPlugin '%s' failed on %s: %s",
                        type(tp_plugin).__name__,
                        file_path.name,
                        exc,
                    )

        file_hash = hashlib.sha256(text.encode()).hexdigest()

        existing = await self._doc_repo.get_by_path(str(file_path))
        if existing and existing.file_hash == file_hash:
            # File content is unchanged. If the graph was never built for this document
            # (e.g. a previous run had broken env vars), rebuild it now using the chunks
            # already stored in the database — no need to re-embed or re-chunk.
            if self._graph_state_repo is not None:
                state = await self._graph_state_repo.get_by_document(existing.id)
                if state is None:
                    await self._rebuild_graph_from_db(
                        existing.id, file_hash, file_path.name
                    )
            return False  # content unchanged

        # Remove stale chunks, graph nodes, and graph state before re-indexing.
        if existing:
            await self._chunk_repo.delete_by_document(existing.id)
            await self._graph_service.delete_document(existing.id)
            if self._graph_state_repo is not None:
                await self._graph_state_repo.delete_by_document(existing.id)

        raw_chunks = self._chunker.chunk(text)
        if not raw_chunks:
            return False

        doc_id = str(uuid.uuid4())
        chunks = [
            Chunk(
                id=str(uuid.uuid4()),
                document_id=doc_id,
                chunk_index=i,
                content=content,
                char_start=char_start,
                char_end=char_end,
            )
            for i, (content, char_start, char_end) in enumerate(raw_chunks)
        ]

        chunk_texts = [c.content for c in chunks]
        embeddings = await loop.run_in_executor(
            None, self._embedding_service.generate_batch, chunk_texts
        )

        # Document row must exist before chunks are inserted (FK constraint).
        doc = IndexedDocument(
            id=doc_id,
            workspace_path=workspace_path,
            file_path=str(file_path),
            file_hash=file_hash,
            char_count=len(text),
            chunk_count=len(chunks),
            indexed_at=datetime.now(UTC).isoformat(),
        )
        await self._doc_repo.upsert(doc)

        await self._chunk_repo.save_batch(chunks, embeddings)

        # Abbreviation extraction is best-effort — failure must not abort indexing.
        await self._extract_and_save_abbreviations(doc_id, text, file_path.name)

        # Graph build — deferred to Pass 2 during workspace indexing (unblocks search
        # sooner), or run inline for single-file indexing (watcher-triggered).
        if deferred is not None:
            deferred.append((doc_id, file_hash, chunk_texts, file_path.name))
        else:
            await self._build_graph_incremental(doc_id, file_hash, chunks, file_path.name)

        logger.debug("Indexed %s (%d chunks)", file_path.name, len(chunks))
        return True

    async def _run_graph_pass(
        self,
        builds: list[tuple[str, str, list[str], str]],
    ) -> None:
        """Pass 2: build knowledge graph for documents indexed in Pass 1.

        Runs as a background asyncio task after index_workspace returns so
        the user can search immediately while graph extraction proceeds.
        Each build is best-effort — failures are logged and do not block
        the remaining documents.
        """
        for doc_id, file_hash, chunk_texts, file_name in builds:
            await asyncio.sleep(0)  # yield between documents; allows cancellation
            try:
                if self._graph_state_repo is not None:
                    state = await self._graph_state_repo.get_by_document(doc_id)
                    if state is not None and state.file_hash == file_hash:
                        continue  # already built (e.g. unchanged re-index)
                await self._graph_service.build_from_chunks(doc_id, chunk_texts)
                if self._graph_state_repo is not None:
                    await self._graph_state_repo.save(doc_id, file_hash)
                logger.debug("[PASS 2] Graph built for %s", file_name)
            except Exception as exc:
                logger.warning("[PASS 2] Graph build failed for %s: %s", file_name, exc)
        logger.info("[PASS 2] Graph extraction complete for %d document(s)", len(builds))

        # Pass 3 — cross-document entity linking. Runs once after all documents
        # in this batch are graphed so shared canonical names across the full
        # workspace are connected. No LLM calls; pure SQL.
        try:
            new_edges = await self._graph_service.link_shared_entities()
            if new_edges:
                logger.info("[PASS 3] Cross-document linking: %d new SIMILAR_TO edge(s)", new_edges)
        except Exception as exc:
            logger.warning("[PASS 3] Cross-document linking failed: %s", exc)

    async def _extract_and_save_abbreviations(
        self,
        doc_id: str,
        text: str,
        file_name: str,
    ) -> None:
        """Extract abbreviations from *text* and persist them for *doc_id*.

        Clears any previously stored abbreviations for the document first so
        re-indexing does not accumulate stale entries.  Failures are logged
        and swallowed — abbreviation extraction is never a reason to fail
        indexing.
        """
        if self._abbreviation_extractor is None or self._abbreviation_repo is None:
            return
        try:
            await self._abbreviation_repo.delete_by_document(doc_id)
            matches = self._abbreviation_extractor.extract(text)
            if matches:
                await self._abbreviation_repo.save_batch(doc_id, matches)
                logger.debug(
                    "Extracted %d abbreviation(s) from %s", len(matches), file_name
                )
        except Exception as exc:
            logger.warning(
                "Abbreviation extraction failed for %s: %s", file_name, exc
            )

    async def _build_graph_incremental(
        self,
        doc_id: str,
        file_hash: str,
        chunks: list,
        file_name: str,
    ) -> None:
        """Build the knowledge graph for doc_id, skipping when the hash is unchanged.

        Checks GraphStateRepository first — if the stored hash matches the current
        file hash the graph is already up to date and the (expensive) LLM extraction
        calls are skipped.  On success the new hash is persisted so future runs
        benefit from the same shortcut.

        Failures are logged and swallowed — graph build must never abort indexing.
        """
        try:
            if self._graph_state_repo is not None:
                state = await self._graph_state_repo.get_by_document(doc_id)
                if state is not None and state.file_hash == file_hash:
                    logger.debug(
                        "Graph build skipped for %s — content unchanged", file_name
                    )
                    return

            await self._graph_service.build_from_chunks(
                doc_id,
                [c.content for c in chunks],
            )

            if self._graph_state_repo is not None:
                await self._graph_state_repo.save(doc_id, file_hash)

        except Exception as exc:
            logger.warning("Graph build failed for %s: %s", file_name, exc)

    async def _rebuild_graph_from_db(
        self,
        doc_id: str,
        file_hash: str,
        file_name: str,
    ) -> None:
        """Rebuild the knowledge graph for doc_id using chunks already in the database.

        Called when a file's content is unchanged but graph_state is missing — this
        happens when a previous indexing run completed successfully for chunks/embeddings
        but failed during graph extraction (e.g. LLM env vars not set).

        Loads chunk text from the chunks table and delegates to _build_graph_incremental.
        Failures are logged and swallowed.
        """
        try:
            async with self._chunk_repo._conn.execute(
                "SELECT content FROM chunks WHERE document_id = ? ORDER BY chunk_index",
                (doc_id,),
            ) as cur:
                rows = await cur.fetchall()

            if not rows:
                logger.debug("No chunks found in DB for %s — skipping graph rebuild", file_name)
                return

            chunk_texts = [row[0] for row in rows]
            logger.info(
                "Rebuilding graph for %s from %d stored chunks", file_name, len(chunk_texts)
            )

            await self._graph_service.build_from_chunks(doc_id, chunk_texts)

            if self._graph_state_repo is not None:
                await self._graph_state_repo.save(doc_id, file_hash)

        except Exception as exc:
            logger.warning("Graph rebuild from DB failed for %s: %s", file_name, exc)
